"""Document upload ingester.

Parses uploaded files (.md, .txt, .pdf, .docx) into IngestedFile objects
that feed into the same extraction -> synthesis -> write pipeline as code.
"""

import os
from pipeline.ingest.models import IngestedFile
import structlog

log = structlog.get_logger()

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}
MAX_DOC_SIZE_BYTES = 2 * 1024 * 1024  # 2MB per doc
MAX_EXTRACTED_CHARS = 100_000  # cap text extraction to ~25K tokens


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    import pdfplumber
    import io

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from .docx using python-docx."""
    from docx import Document
    import io

    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def ingest_uploaded_doc(
    filename: str,
    file_bytes: bytes,
) -> IngestedFile | None:
    """Parse a single uploaded document into an IngestedFile.

    Returns None if the file can't be parsed or is empty.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        log.warning("upload.ingest.unsupported", filename=filename, ext=ext)
        return None

    if len(file_bytes) > MAX_DOC_SIZE_BYTES:
        log.warning("upload.ingest.too_large", filename=filename,
                    size=len(file_bytes), limit=MAX_DOC_SIZE_BYTES)
        return None

    try:
        if ext in (".md", ".txt"):
            content = file_bytes.decode("utf-8", errors="replace")
        elif ext == ".pdf":
            content = _extract_pdf(file_bytes)
        elif ext == ".docx":
            content = _extract_docx(file_bytes)
        else:
            return None
    except Exception as e:
        log.error("upload.ingest.parse_error", filename=filename, error=str(e)[:200])
        return None

    if not content or not content.strip():
        log.warning("upload.ingest.empty", filename=filename)
        return None

    # Truncate to prevent enormous docs from blowing up extraction
    content = content[:MAX_EXTRACTED_CHARS]

    language = "markdown" if ext == ".md" else "text"

    log.info("upload.ingest.success", filename=filename,
             chars=len(content), ext=ext)

    return IngestedFile(
        path=f"upload://{filename}",
        content=content,
        language=language,
        size_bytes=len(file_bytes),
        last_modified="",
        source_type="upload",
    )
